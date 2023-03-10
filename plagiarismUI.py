import tkinter as tk
from tkinter import ttk
from tkinter import filedialog, Text
from tkinter import messagebox
from tkinter import *
import docx
import fuzzyMatch as fsm

root = tk.Tk()
root.resizable(0,0)
root.title("Plagiarism Checker")
root.geometry("1250x625")

#threshold = 0.5

def browse_file1():
    global contents1
    file_path = filedialog.askopenfilename()
    lblfn1.config(text=file_path)
    document = docx.Document(file_path)
    contents1 = "\n".join([paragraph.text for paragraph in document.paragraphs])
    txtf1.delete("1.0", "end")
    txtf1.insert("1.0", contents1)

def browse_file2():
    global contents2
    file_path = filedialog.askopenfilename()
    lblfn2.config(text=file_path)
    document = docx.Document(file_path)
    contents2 = "\n".join([paragraph.text for paragraph in document.paragraphs])
    txtf2.delete("1.0", "end")
    txtf2.insert("1.0", contents2)

def check_plagiarism():
    threshold = slider.get()
    
    file1 = lblfn1["text"]
    file2 = lblfn2["text"]

    if not file1 or not file2:
        messagebox.showerror("Error", "Please select two files first")
        return

    #messagebox.showinfo("Threshold", threshold)
    #Total # of String
    n = min(len(contents1.split()), len(contents2.split()))

    # Add your plagiarism checking code here
    similarity = fsm.compare_documents(contents1, contents2, threshold)

    messagebox.showinfo("Plagiarism Checker Result", "Match: {} \nTotal Number of Strings: {}\n Similarity Results: {:.2f}%".format(similarity, n, float(similarity/n)*100))


#Label for Title
lbltitle = tk.Label(root, text="PLAGIARISM CHECKER", font=("Arial", 14), pady=5)
lbltitle.pack(pady=5, fill="x")
lbltitle.config(anchor="center")

#Frame for file insert 1
frmMid = tk.Frame(root, width=350, height=500, bg='white', highlightthickness=2, highlightbackground="black")
frmMid.pack(fill="both", expand = True, padx=10, pady=20)

#Label for Inserting file 1
lblf1 = tk.Label(frmMid, text="Insert File 1", font=("Arial", 12), pady=10, bg="white")
lblf1.grid(row=0, column=1, sticky="w")

#Label for file 1 name
lblfn1 = tk.Label(frmMid, text="", font=("Arial", 10), pady=10, bg="white")
lblfn1.grid(row=1, column=1, sticky="w")

#Textfield for file 1 content
txtf1 = tk.Text(frmMid, font=("Arial", 10), height=18, width=70, highlightthickness=2, highlightbackground="black")
txtf1.grid (row=2, column=1, sticky="w")

#Button for inserting file 1
btnf1 = tk.Button(frmMid, text="Insert File 1", font=("Arial", 10), pady=3, bg="white", command=browse_file1)
btnf1.grid(row=4, column=1, sticky="w")

#Label for inserting file 2
lblf2 = tk.Label(frmMid, text="Insert File 2", font=("Arial", 12), pady=10, bg="white")
lblf2.grid(row=0, column=6, sticky="w")

#Label for file 1 name
lblfn2 = tk.Label(frmMid, text="", font=("Arial", 10), pady=10, bg="white")
lblfn2.grid(row=1, column=6, sticky="w")

#Textfield for file 2 content
txtf2 = tk.Text(frmMid, font=("Arial", 10), height=18, width=70, highlightthickness=2, highlightbackground="black")
txtf2.grid (row=2, column=6, sticky="w")

#Button for inserting file 2
btnf2 = tk.Button(frmMid, text="Insert File 2", font=("Arial", 10), pady=3, bg="white", command=browse_file2)
btnf2.grid(row=4, column=6, sticky="w")

#Treshold
slider = tk.Scale(frmMid, from_=0, to=1, font=("Arial", 11), bg="white", orient=tk.HORIZONTAL, length=150, resolution=0.01)
slider.set(0.5)
slider.grid(row=5, column=3, sticky="w")

#Button for Plagirism Checking
btnc = tk.Button(frmMid, text="Check Plagiarism", font=("Arial", 10), padx=20, pady=3, bg="white", command=check_plagiarism)
btnc.grid(row=7, column=3, sticky="w")

#Blank Spaces
lblb = tk.Label(frmMid, text="       ", bg="white").grid(row=1, column=0, sticky="w")
lblb1 = tk.Label(frmMid, text="      ", bg="white").grid(row=0, column=3, sticky="w")
lblb2 = tk.Label(frmMid, text="      ", bg="white").grid(row=0, column=4, sticky="w")
lblb3 = tk.Label(frmMid, text="      ", bg="white").grid(row=2, column=3, sticky="w")
lblb4 = tk.Label(frmMid, text="      ", bg="white").grid(row=2, column=4, sticky="w")
lblb5 = tk.Label(frmMid, text="      ", bg="white").grid(row=3, column=4, sticky="w")
lblb5 = tk.Label(frmMid, text="      ", bg="white").grid(row=6, column=0, sticky="w")

root.mainloop()
